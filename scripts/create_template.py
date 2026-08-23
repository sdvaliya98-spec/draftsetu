from docx import Document

doc = Document()
doc.add_heading('વેચાણ અઘાટનો દસ્તાવેજ', 0)

doc.add_paragraph('Survey No: {{survey_no}}')
doc.add_paragraph('Area: {{area}}')
doc.add_paragraph('Amount: {{amount}}')

p = doc.add_paragraph('Buyer Details:\n')
p.add_run('Name: {{buyer_name}}\n')
p.add_run('Address: {{buyer_address}}\n')
p.add_run('PAN: {{buyer_pan}}\n')

p2 = doc.add_paragraph('Seller Details:\n')
p2.add_run('Name: {{seller_name}}\n')
p2.add_run('Address: {{seller_address}}\n')
p2.add_run('PAN: {{seller_pan}}\n')

doc.save('backend/templates/master.docx')
print("Template created successfully.")
