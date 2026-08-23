"""
DOCX Template Engine — Canonical Document Generation Service
============================================================
Architecture: DOCX Template → Variable Fill (docxtpl/Jinja2) → Output DOCX → [Word COM / LibreOffice → PDF]

PDF Engine Priority:
  1. Microsoft Word COM via docx2pdf (Windows — perfect Gujarati font fidelity)
  2. LibreOffice headless subprocess (cross-platform fallback)

This module is the SINGLE source of truth for document rendering.
NO HTML rendering. NO browser printing. NO CSS layout approximations.
"""

import os
import re
import uuid
import time
import shutil
import logging
import subprocess
import threading
from typing import Optional

from docx import Document
from docxtpl import DocxTemplate
from backend.core.config import settings

logger = logging.getLogger("backend.docx_engine")

# Concurrency protection: limit simultaneous heavy renders
RENDER_LOCK = threading.Semaphore(3)


# ─── VARIABLE EXTRACTION ────────────────────────────────────────────────────

def extract_variables_from_docx(file_path: str) -> dict:
    """
    Extracts all variables and Jinja2 loops from a .docx file.
    Scans paragraphs, tables (recursive), headers, and footers.
    Supports repeater block tags: {% for x in X %} and {% endfor %}.
    Returns a dictionary containing "groups" and "single_variables".
    """
    loop_pattern = re.compile(r'{%\s*for\s+(\w+)\s+in\s+(\w+)\s*%}')
    var_pattern = re.compile(r'\{\{([^}]+)\}\}')

    all_texts: list[str] = []

    logger.info(f"🔍 EXTRACTING VARIABLES: {os.path.basename(file_path)}")

    def scan_text(text: str):
        if text:
            all_texts.append(text)

    def scan_paragraph(p):
        scan_text(p.text)

    def scan_table(table):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    scan_paragraph(p)
                for nested in cell.tables:
                    scan_table(nested)

    try:
        if not os.path.exists(file_path):
            logger.error(f"❌ File not found: {file_path}")
            return {"groups": {}, "single_variables": []}

        doc = Document(file_path)
        logger.info(f"✅ Opened DOCX: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

        # Scan headers
        for section in doc.sections:
            try:
                for p in section.header.paragraphs:
                    scan_paragraph(p)
                for t in section.header.tables:
                    scan_table(t)
            except Exception as e:
                logger.debug(f"Header scan skip: {e}")

        # Scan main body
        for p in doc.paragraphs:
            scan_paragraph(p)
        for t in doc.tables:
            scan_table(t)

        # Scan footers
        for section in doc.sections:
            try:
                for p in section.footer.paragraphs:
                    scan_paragraph(p)
                for t in section.footer.tables:
                    scan_table(t)
            except Exception as e:
                logger.debug(f"Footer scan skip: {e}")

        # Pass 1: Scan all text for loop blocks
        iterators = {}
        detected_groups = []
        detected_groups_set = set()
        for text in all_texts:
            for m in loop_pattern.finditer(text):
                iterator = m.group(1).strip()
                group = m.group(2).strip()
                iterators[iterator] = group
                if group not in detected_groups_set:
                    detected_groups_set.add(group)
                    detected_groups.append(group)
                    logger.info(f"[LOOP DETECTED] {group}")

        # Pass 2: Parse variables
        groups = {g: [] for g in detected_groups}
        groups_seen = {g: set() for g in detected_groups}
        single_variables = []
        single_variables_set = set()

        for text in all_texts:
            for m in var_pattern.finditer(text):
                var_content = m.group(1).strip()
                if '.' in var_content:
                    parts = var_content.split('.', 1)
                    prefix = parts[0].strip()
                    field_name = parts[1].strip()

                    if prefix in iterators:
                        g = iterators[prefix]
                        if field_name not in groups_seen[g]:
                            groups_seen[g].add(field_name)
                            groups[g].append(field_name)
                    else:
                        if var_content not in single_variables_set:
                            single_variables_set.add(var_content)
                            single_variables.append(var_content)
                else:
                    if var_content not in iterators:
                        if var_content not in single_variables_set:
                            single_variables_set.add(var_content)
                            single_variables.append(var_content)

        result = {
            "groups": groups,
            "single_variables": single_variables
        }

        logger.info(f"[EXTRACT] Found loop groups: {list(result['groups'].keys())}, single variables: {result['single_variables']}")
        return result

    except Exception as e:
        logger.critical(f"🔥 EXTRACTION FATAL: {e}", exc_info=True)
        return {"groups": {}, "single_variables": []}


# ─── DOCX RENDERING ─────────────────────────────────────────────────────────

def _flatten_heirs(heirs_list: list) -> list:
    """Recursively flatten heirs that may contain nested 'children' arrays.
    Each heir dict is expected to have an optional 'index' field representing its hierarchical position.
    The function returns a flat list of heir dictionaries preserving the original index hierarchy (e.g., parent 1 -> child 1.1).
    """
    flat = []
    def _walk(item, parent_index=None, local_idx=1):
        if parent_index:
            combined = f"{parent_index}.{local_idx}"
        else:
            combined = str(local_idx)
        # Copy without children
        flat_item = {k: v for k, v in item.items() if k != 'children'}
        flat_item['index'] = combined
        flat.append(flat_item)
        children = item.get('children')
        if isinstance(children, list):
            for i, child in enumerate(children, start=1):
                _walk(child, combined, i)
    for i, heir in enumerate(heirs_list, start=1):
        _walk(heir, parent_index=None, local_idx=i)
    return flat

def _normalize_context(data: dict) -> dict:
    """
    Converts flat/nested data dict to docxtpl-compatible context.
    Handles nested repeaters and recursive lists/dicts.
    Ensures all primitive values are formatted strings (no None values),
    and automatically formats ISO date strings (YYYY-MM-DD) to Indian format (DD/MM/YYYY).
    """
    from backend.utils.date_utils import formatDateForDocument

    def normalize_val(val):
        if val is None:
            return ""
        if isinstance(val, list):
            return [normalize_val(item) for item in val]
        if isinstance(val, dict):
            return {fk: normalize_val(fv) for fk, fv in val.items()}
        # Primitive value
        return formatDateForDocument(str(val))

    # Perform recursive normalization
    normalized = normalize_val(data)
    # Flatten hierarchical lists (like HEIRS, family_members, etc.) if present
    if isinstance(normalized, dict):
        heir_keys = ['heirs', 'family_members', 'members', 'heir_tree']
        for k, v in list(normalized.items()):
            if k.lower() in heir_keys and isinstance(v, list) and len(v) > 0:
                if any(isinstance(item, dict) and 'children' in item for item in v):
                    try:
                        flattened = _flatten_heirs(v)
                        normalized[k] = flattened
                        logger.info(f"[{k} FLATTEN] Flattened hierarchical list '{k}', count: {len(flattened)}")
                        import json
                        logger.info(f"[{k} FLATTENED RESULT]: {json.dumps(flattened, indent=2)}")
                    except Exception as e:
                        logger.error(f"Failed to flatten hierarchical list '{k}': {e}")
    if isinstance(normalized, dict):
        return normalized
    return {}


def render_docx_template(
    template_path: str,
    data: dict,
    output_path: str,
    tracking_id: Optional[str] = None,
) -> str:
    """
    Renders a DOCX template with user data using docxtpl (Jinja2 engine).

    Args:
        template_path: Absolute path to the .docx template file.
        data: Dictionary of variable values (can include lists for repeaters).
        output_path: Full path where the rendered .docx will be saved.
        tracking_id: Optional tracking ID for logging.

    Returns:
        output_path if successful.

    Raises:
        FileNotFoundError if template is missing.
        Exception if docxtpl render fails.
    """
    tid = tracking_id or uuid.uuid4().hex[:8]
    start = time.perf_counter()

    logger.info(f"📄 DOCX RENDER START [{tid}]: template={os.path.basename(template_path)}")

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    with RENDER_LOCK:
        try:
            doc = DocxTemplate(template_path)
            context = _normalize_context(data)
            # Log final HEIRS array for debugging
            if 'HEIRS' in context:
                logger.info(f"[RENDER {tid}] Final HEIRS payload: {context['HEIRS']}")
            logger.info(f"[RENDER {tid}] Context keys: {list(context.keys())}")

            # Compatibility mapping for paragraph repeater (single textarea to universal paragraphs list)
            text_value = context.get("EXTRA_PARAGRAPHS_TEXT", "")
            if text_value and str(text_value).strip():
                # Normalize line endings to standard Unix \n
                normalized_text = str(text_value).replace("\r\n", "\n").replace("\r", "\n")
                paragraphs = [
                    {
                        "text": p.strip()
                    }
                    for p in re.split(r"\n\s*\n", normalized_text)
                    if p.strip()
                ]
                context["EXTRA_PARAGRAPHS"] = paragraphs
            elif "para.text" in context:
                para_value = context.get("para.text")
                if para_value and str(para_value).strip():
                    normalized_text = str(para_value).replace("\r\n", "\n").replace("\r", "\n")
                    paragraphs = [
                        {
                            "text": p.strip()
                        }
                        for p in re.split(r"\n\s*\n", normalized_text)
                        if p.strip()
                    ]
                    context["EXTRA_PARAGRAPHS"] = paragraphs
                else:
                    context["EXTRA_PARAGRAPHS"] = []
            elif "EXTRA_PARAGRAPHS" in context:
                # Keep and normalize pre-existing EXTRA_PARAGRAPHS
                existing_paras = context["EXTRA_PARAGRAPHS"]
                if isinstance(existing_paras, list):
                    normalized_paras = []
                    for item in existing_paras:
                        if isinstance(item, dict) and "text" in item:
                            text_val = item["text"]
                            if text_val is not None:
                                text_str = str(text_val).replace("\r\n", "\n").replace("\r", "\n")
                                item["text"] = text_str
                            normalized_paras.append(item)
                        elif isinstance(item, str):
                            normalized_paras.append({
                                "text": item.replace("\r\n", "\n").replace("\r", "\n")
                            })
                        else:
                            normalized_paras.append(item)
                    context["EXTRA_PARAGRAPHS"] = normalized_paras
            else:
                context["EXTRA_PARAGRAPHS"] = []

            context.setdefault("EXTRA_PARAGRAPHS", [])
            logger.info(
                f"EXTRA_PARAGRAPHS Generated Count: {len(context.get('EXTRA_PARAGRAPHS', []))}"
            )

            doc.render(context)
            doc.save(output_path)

            duration = time.perf_counter() - start
            logger.info(f"✅ DOCX RENDER COMPLETE [{tid}]: {os.path.basename(output_path)} in {duration:.3f}s")
            return output_path

        except Exception as e:
            logger.error(f"❌ DOCX RENDER FAILED [{tid}]: {e}", exc_info=True)
            raise


# ─── PDF CONVERSION ──────────────────────────────────────────────────────────

def _find_libreoffice() -> Optional[str]:
    """
    Locate the LibreOffice soffice binary on Windows, Linux, and macOS.
    Strategy:
      1. Hard-coded common Windows install paths (C, D, E drives)
      2. Windows Registry HKLM lookup (catches non-standard installs)
      3. shutil.which() for PATH-based installs and Linux/Mac
    """
    candidates = [
        # Windows — standard 64-bit
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files\LibreOffice 7\program\soffice.exe",
        r"C:\Program Files\LibreOffice 6\program\soffice.exe",
        # Windows — 32-bit on 64-bit OS
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice 7\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice 6\program\soffice.exe",
        # Alternate drive installs
        r"D:\LibreOffice\program\soffice.exe",
        r"D:\Program Files\LibreOffice\program\soffice.exe",
        r"E:\LibreOffice\program\soffice.exe",
        # Linux
        "/usr/bin/soffice",
        "/usr/bin/libreoffice",
        "/usr/local/bin/soffice",
        "/usr/local/bin/libreoffice",
        "/snap/bin/libreoffice",
        # macOS
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ]

    for path in candidates:
        if os.path.exists(path):
            logger.info(f"✅ [LIBREOFFICE] Found via candidate list: {path}")
            return path

    # Windows Registry lookup (catches non-standard install directories)
    try:
        import winreg
        for root_key in (winreg.HKEY_LOCAL_MACHINE, winreg.HKEY_CURRENT_USER):
            for sub in (
                r"SOFTWARE\LibreOffice\UNO\InstallPath",
                r"SOFTWARE\WOW6432Node\LibreOffice\UNO\InstallPath",
            ):
                try:
                    with winreg.OpenKey(root_key, sub) as k:
                        install_path, _ = winreg.QueryValueEx(k, "")
                        candidate = os.path.join(install_path, "soffice.exe")
                        if os.path.exists(candidate):
                            logger.info(f"✅ [LIBREOFFICE] Found via Windows Registry: {candidate}")
                            return candidate
                except FileNotFoundError:
                    pass
    except ImportError:
        pass  # Not on Windows — skip registry

    # Fall back to PATH / shutil.which
    for cmd in ("soffice", "libreoffice"):
        found = shutil.which(cmd)
        if found:
            logger.info(f"✅ [LIBREOFFICE] Found via shutil.which('{cmd}'): {found}")
            return found

    logger.info("ℹ️ [LIBREOFFICE] Not found — will use Microsoft Word (docx2pdf) if available.")
    return None


def _check_docx2pdf() -> bool:
    """
    Check whether docx2pdf (Microsoft Word COM automation) is available.
    Works on Windows when Microsoft Word is installed.
    Returns True only when Word COM dispatch succeeds.
    """
    if os.name != "nt":
        return False  # docx2pdf's Word COM path is Windows-only

    try:
        import docx2pdf  # noqa: F401
        import win32com.client
        import pythoncom
        
        pythoncom.CoInitialize()
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Quit()
        finally:
            pythoncom.CoUninitialize()
            
        logger.info("✅ [DOCX2PDF] Microsoft Word COM available — PDF engine ready.")
        return True
    except ImportError:
        logger.warning("⚠️ [DOCX2PDF] Package not installed. Run: pip install docx2pdf")
        return False
    except Exception as e:
        logger.warning(f"⚠️ [DOCX2PDF] Word COM check failed: {e}")
        return False


# ── Detect available PDF engines at module load ───────────────────────────────

import shutil

LIBREOFFICE_PATH = None
LIBREOFFICE_AVAILABLE = False

WINDOWS_LIBREOFFICE_PATHS = [
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    r"D:\Program Files\LibreOffice\program\soffice.exe",
    r"D:\LibreOffice\program\soffice.exe",
]

for path in WINDOWS_LIBREOFFICE_PATHS:
    normalized = os.path.normpath(path)
    if os.path.isfile(normalized):
        LIBREOFFICE_PATH = normalized
        LIBREOFFICE_AVAILABLE = True
        break

if not LIBREOFFICE_AVAILABLE:
    LIBREOFFICE_PATH = _find_libreoffice()
    if LIBREOFFICE_PATH:
        LIBREOFFICE_PATH = os.path.normpath(LIBREOFFICE_PATH)
        LIBREOFFICE_AVAILABLE = os.path.isfile(LIBREOFFICE_PATH)

print("LIBREOFFICE AVAILABLE:", LIBREOFFICE_AVAILABLE)
print("LIBREOFFICE PATH:", LIBREOFFICE_PATH)

DOCX2PDF_AVAILABLE: bool = _check_docx2pdf()
PDF_ENGINE_AVAILABLE: bool = DOCX2PDF_AVAILABLE or LIBREOFFICE_AVAILABLE

def _safe_print(msg: str):
    try:
        print(msg)
    except UnicodeEncodeError:
        try:
            print(msg.encode('ascii', errors='backslashreplace').decode('ascii'))
        except Exception:
            pass

# Startup diagnostics — printed in the backend console/log on every restart
_safe_print(f"[PDF ENGINE] Microsoft Word (docx2pdf) : {'AVAILABLE' if DOCX2PDF_AVAILABLE else 'not found'}")
_safe_print(f"[PDF ENGINE] LibreOffice               : {LIBREOFFICE_PATH if LIBREOFFICE_AVAILABLE else 'not found'}")
_safe_print(f"[PDF ENGINE] PDF export                : {'ENABLED' if PDF_ENGINE_AVAILABLE else 'DISABLED'}")
logger.info(
    f"[PDF ENGINE] docx2pdf={DOCX2PDF_AVAILABLE} | "
    f"libreoffice={LIBREOFFICE_AVAILABLE} | "
    f"pdf_available={PDF_ENGINE_AVAILABLE}"
)


word_pdf_lock = threading.Lock()


def kill_zombie_winword():
    """
    Terminate orphan WINWORD.EXE processes older than 2 minutes using psutil.
    """
    try:
        import psutil
        current_time = time.time()
        for proc in psutil.process_iter(['pid', 'name', 'create_time']):
            try:
                if proc.info['name'] and proc.info['name'].upper() == 'WINWORD.EXE':
                    age_seconds = current_time - proc.info['create_time']
                    if age_seconds > 120:
                        logger.warning(f"Killing zombie WINWORD.EXE with PID {proc.info['pid']} (age: {age_seconds:.1f}s)")
                        proc.terminate()
                        try:
                            proc.wait(timeout=2)
                        except psutil.TimeoutExpired:
                            proc.kill()
            except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
                pass
    except Exception as e:
        logger.error(f"Failed to run zombie winword killer: {e}")


def _convert_via_word(docx_path: str, output_dir: str) -> str:
    """
    Convert DOCX → PDF using Microsoft Word COM automation (Windows only).
    Preserves ALL Word formatting, Gujarati fonts, and complex layouts exactly
    because it uses Word's own rendering engine.

    Raises RuntimeError on failure.
    """
    import pythoncom
    import win32com.client

    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    pdf_path = os.path.join(output_dir, f"{base_name}.pdf")

    logger.info("💾 SAVING PDF")
    _safe_print(f"[WORD COM] {docx_path} -> {pdf_path}")

    start = time.perf_counter()

    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0  # wdAlertsNone = 0

        abs_docx = os.path.abspath(docx_path)
        abs_pdf  = os.path.abspath(pdf_path)

        # Open with explicit parameters to avoid COM returning method object
        # Parameters: FileName, ConfirmConversions, ReadOnly, AddToRecentFiles,
        #             PasswordDocument, PasswordTemplate, Revert, WritePasswordDocument,
        #             WritePasswordTemplate, Format
        doc = word.Documents.Open(
            abs_docx,   # FileName
            False,      # ConfirmConversions
            False,      # ReadOnly (must be False to allow SaveAs)
            False,      # AddToRecentFiles
        )

        if doc is None:
            raise RuntimeError("Word.Documents.Open returned None — file may be locked or corrupted.")

        # wdFormatPDF = 17
        # Use SaveAs2 when available (Word 2010+), fall back to SaveAs
        try:
            doc.SaveAs2(abs_pdf, FileFormat=17)
        except AttributeError:
            doc.SaveAs(abs_pdf, FileFormat=17)

    except Exception as e:
        logger.error("⚠️ WORD COM FAILED")
        raise RuntimeError(f"Microsoft Word PDF conversion failed: {e}") from e
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception as ex:
                logger.warning(f"Error closing doc: {ex}")
        if word is not None:
            try:
                word.Quit()
            except Exception as ex:
                logger.warning(f"Error quitting word: {ex}")
        pythoncom.CoUninitialize()

    if not os.path.exists(pdf_path):
        raise RuntimeError(
            f"Word COM conversion appeared to succeed but PDF not found at: {pdf_path}"
        )

    duration = time.perf_counter() - start
    logger.info("✅ PDF READY")
    logger.info(f"[WORD COM] PDF created in {duration:.3f}s: {os.path.basename(pdf_path)}")
    return pdf_path


# ── Engine 2: LibreOffice subprocess ─────────────────────────────────────────

def _convert_via_libreoffice(docx_path: str, output_dir: str) -> str:
    """
    Convert DOCX → PDF using LibreOffice headless subprocess.
    Used as fallback when Microsoft Word is not available.

    Raises Exception on failure.
    """
    from pathlib import Path

    global LIBREOFFICE_PATH
    if LIBREOFFICE_PATH:
        LIBREOFFICE_PATH = os.path.normpath(LIBREOFFICE_PATH)

    if not LIBREOFFICE_PATH or not os.path.isfile(LIBREOFFICE_PATH):
        raise Exception(
            f"Invalid LibreOffice executable: {LIBREOFFICE_PATH}"
        )

    out_dir = os.path.normpath(output_dir)
    docx_path = os.path.normpath(docx_path)

    # 1. Setup isolated user profile directory (cross-platform safe)
    user_profile_uuid = uuid.uuid4().hex
    if os.name == 'nt':
        temp_dir_raw = os.environ.get('TEMP', 'C:\\Temp')
        profile_disk_path = os.path.join(temp_dir_raw, f"libreoffice_user_{user_profile_uuid}")
        temp_dir_uri = temp_dir_raw.replace('\\', '/').lstrip('/')
        user_install_uri = f"file:///{temp_dir_uri}/libreoffice_user_{user_profile_uuid}"
    else:
        profile_disk_path = f"/tmp/libreoffice_user_{user_profile_uuid}"
        user_install_uri = f"file:///tmp/libreoffice_user_{user_profile_uuid}"

    command = [
        str(LIBREOFFICE_PATH),
        f"-env:UserInstallation={user_install_uri}",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(out_dir),
        str(docx_path)
    ]

    print("RUNNING COMMAND:", command)

    try:
        result = subprocess.run(
            command,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            shell=False
        )

        print("STDOUT:", result.stdout)
        print("STDERR:", result.stderr)
        print("RETURN CODE:", result.returncode)

        if result.returncode != 0:
            raise Exception(
                f"LibreOffice PDF conversion failed: {result.stderr}"
            )

        pdf_path = os.path.join(
            out_dir,
            Path(docx_path).stem + ".pdf"
        )

        if not os.path.exists(pdf_path):
            raise Exception(
                f"PDF file not generated: {pdf_path}"
            )

        return pdf_path
    finally:
        # 2. Guarantee cleanup of the temporary profile directory to avoid disk space/inode exhaustion
        if os.path.exists(profile_disk_path):
            try:
                shutil.rmtree(profile_disk_path)
                logger.info(f"🧹 Cleaned up temporary LibreOffice profile: {profile_disk_path}")
            except Exception as cleanup_err:
                logger.warning(f"Failed to clean up LibreOffice profile directory {profile_disk_path}: {cleanup_err}")


# ── Public API ────────────────────────────────────────────────────────────────

def convert_docx_to_pdf(docx_path: str, output_dir: Optional[str] = None) -> str:
    """
    Convert a rendered DOCX file to PDF using the best available engine:
      1. Microsoft Word COM (docx2pdf) — preferred on Windows, pixel-perfect fidelity
      2. LibreOffice headless subprocess — cross-platform fallback

    Args:
        docx_path:   Absolute path to the .docx file to convert.
        output_dir:  Directory to write the PDF. Defaults to same dir as docx.

    Returns:
        Absolute path to the generated .pdf file.

    Raises:
        RuntimeError if no PDF engine is available or conversion fails.
    """
    if not PDF_ENGINE_AVAILABLE:
        raise RuntimeError(
            "No PDF engine available. "
            "On Windows: Microsoft Word must be installed (and pip install docx2pdf). "
            "On Linux/macOS: install LibreOffice (sudo apt install libreoffice)."
        )

    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    out_dir = output_dir or os.path.dirname(docx_path)
    os.makedirs(out_dir, exist_ok=True)

    # Engine 1 — Microsoft Word COM (Windows, best quality for Gujarati fonts)
    if DOCX2PDF_AVAILABLE:
        logger.info("🟦 STARTING PDF CONVERSION")
        logger.info("🔒 WAITING FOR WORD LOCK")
        with word_pdf_lock:
            for attempt in range(3):
                try:
                    kill_zombie_winword()
                    return _convert_via_word(docx_path, out_dir)
                except Exception as e:
                    logger.warning(f"Word COM attempt {attempt + 1} failed: {e}")
                    if attempt < 2:
                        logger.info("♻️ RETRYING")
                        time.sleep(1)
                    else:
                        logger.warning(f"⚠️ Word COM failed, trying LibreOffice fallback: {e}")
                        if not LIBREOFFICE_AVAILABLE:
                            raise  # No fallback — surface the error

    # Engine 2 — LibreOffice (fallback / Linux / macOS)
    return _convert_via_libreoffice(docx_path, out_dir)


def libreoffice_available() -> bool:
    """
    Returns True if ANY PDF engine is available (Word COM or LibreOffice).
    Named for backward compatibility — all callers work without changes.
    """
    return PDF_ENGINE_AVAILABLE
